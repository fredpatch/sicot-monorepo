import os
import io
import re
import tempfile
import subprocess
from pathlib import Path

from flask import Flask, request, jsonify
from waitress import serve
import pdfplumber
import pytesseract
from pdf2image import convert_from_bytes
from docx import Document
from langdetect import detect, LangDetectException
from PIL import Image
import openpyxl
import xlrd

# ── Configuration ─────────────────────────────────────────────────────────
app = Flask(__name__)

PORT = int(os.environ.get("OCR_PORT", 5001))
TESSERACT_CMD = os.environ.get(
    "TESSERACT_CMD",
    r"C:\Users\Prime Daily\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"
)
LIBREOFFICE_CMD = os.environ.get(
    "LIBREOFFICE_CMD",
    r"C:\Program Files\LibreOffice\program\soffice.exe"
)

pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

# ── Nettoyage du texte extrait ────────────────────────────────────────────
def nettoyer_texte(texte: str) -> str:
    if not texte:
        return ""

    # Supprimer espaces parasites autour des apostrophes
    # Identifié lors des tests LibreTranslate : "l ' annexe" → "l'annexe"
    texte = re.sub(r"\s+'\s+", "'", texte)

    # Supprimer espaces multiples
    texte = re.sub(r" {2,}", " ", texte)

    # Supprimer lignes vides multiples
    texte = re.sub(r"\n{3,}", "\n\n", texte)

    # Trim
    return texte.strip()

# ── Détection de langue ───────────────────────────────────────────────────
def detecter_langue(texte: str) -> str:
    try:
        if len(texte.strip()) < 20:
            return "inconnu"
        langue = detect(texte)
        # Normaliser vers nos codes supportés
        if langue == "fr":
            return "fr"
        elif langue == "en":
            return "en"
        return langue
    except LangDetectException:
        return "inconnu"

# ── Extracteurs par format ────────────────────────────────────────────────

def extraire_txt(contenu: bytes) -> str:
    """Fichier texte brut — décodage direct."""
    for encoding in ["utf-8", "latin-1", "cp1252"]:
        try:
            return contenu.decode(encoding)
        except UnicodeDecodeError:
            continue
    return contenu.decode("utf-8", errors="replace")


def extraire_pdf(contenu: bytes) -> str:
    """
    PDF natif → pdfplumber
    PDF scanné → pdf2image + Tesseract
    Détection automatique selon présence de texte extractible.
    """
    texte_pages = []

    with pdfplumber.open(io.BytesIO(contenu)) as pdf:
        for i, page in enumerate(pdf.pages):
            texte_page = page.extract_text()

            if texte_page and len(texte_page.strip()) > 20:
                # PDF natif — texte directement extractible
                texte_pages.append(texte_page)
            else:
                # Page vide ou scannée → OCR via Tesseract
                images = convert_from_bytes(
                    contenu,
                    first_page=i + 1,
                    last_page=i + 1,
                    dpi=300
                )
                for image in images:
                    texte_ocr = pytesseract.image_to_string(
                        image,
                        lang="fra+eng",
                        config="--psm 3"
                    )
                    texte_pages.append(texte_ocr)

    return "\n\n".join(texte_pages)


def extraire_docx(contenu: bytes) -> str:
    """Fichier Word .docx — python-docx."""
    doc = Document(io.BytesIO(contenu))
    parties = []

    # Paragraphes principaux
    for para in doc.paragraphs:
        if para.text.strip():
            parties.append(para.text)

    # Texte dans les tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parties.append(cell.text)

    return "\n".join(parties)


def extraire_doc(contenu: bytes) -> str:
    """
    Ancien format Word .doc
    Conversion via LibreOffice headless → .docx → python-docx
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        # Écrire le .doc temporaire
        doc_path = Path(tmp_dir) / "document.doc"
        doc_path.write_bytes(contenu)

        # Convertir en docx via LibreOffice headless
        try:
            subprocess.run(
                [
                    LIBREOFFICE_CMD,
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", tmp_dir,
                    str(doc_path)
                ],
                check=True,
                capture_output=True,
                timeout=30
            )
        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"Échec conversion LibreOffice : {e.stderr.decode()}")
        except FileNotFoundError:
            raise RuntimeError("LibreOffice introuvable — vérifiez LIBREOFFICE_CMD")

        # Lire le .docx généré
        docx_path = Path(tmp_dir) / "document.docx"
        if not docx_path.exists():
            raise RuntimeError("LibreOffice n'a pas généré le fichier .docx")

        return extraire_docx(docx_path.read_bytes())


def extraire_xlsx(contenu: bytes) -> str:
    """Fichier Excel .xlsx — openpyxl."""
    wb = openpyxl.load_workbook(io.BytesIO(contenu), read_only=True, data_only=True)
    parties = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parties.append(f"[Feuille : {sheet_name}]")
        for row in ws.iter_rows(values_only=True):
            valeurs = [str(cell) for cell in row if cell is not None]
            if valeurs:
                parties.append(" | ".join(valeurs))

    return "\n".join(parties)


def extraire_xls(contenu: bytes) -> str:
    """Ancien format Excel .xls — xlrd."""
    wb = xlrd.open_workbook(file_contents=contenu)
    parties = []

    for sheet in wb.sheets():
        parties.append(f"[Feuille : {sheet.name}]")
        for row_idx in range(sheet.nrows):
            valeurs = [
                str(sheet.cell_value(row_idx, col))
                for col in range(sheet.ncols)
                if sheet.cell_value(row_idx, col) != ""
            ]
            if valeurs:
                parties.append(" | ".join(valeurs))

    return "\n".join(parties)


def extraire_image(contenu: bytes) -> str:
    """Image .jpg / .png / .tiff — Tesseract directement."""
    image = Image.open(io.BytesIO(contenu))
    return pytesseract.image_to_string(
        image,
        lang="fra+eng",
        config="--psm 3"
    )


# ── Routeur principal ─────────────────────────────────────────────────────
EXTRACTEURS = {
    ".txt":  extraire_txt,
    ".pdf":  extraire_pdf,
    ".docx": extraire_docx,
    ".doc":  extraire_doc,
    ".xlsx": extraire_xlsx,
    ".xls":  extraire_xls,
    ".jpg":  extraire_image,
    ".jpeg": extraire_image,
    ".png":  extraire_image,
    ".tiff": extraire_image,
    ".tif":  extraire_image,
}


# ── Route principale : POST /extract ─────────────────────────────────────
@app.route("/extract", methods=["POST"])
def extract():
    """
    Reçoit un fichier en multipart/form-data.
    Retourne le texte extrait, la langue détectée et le format.
    """
    if "file" not in request.files:
        return jsonify({"error": "Aucun fichier fourni."}), 400

    fichier = request.files["file"]
    nom = fichier.filename or ""
    extension = Path(nom).suffix.lower()

    if extension not in EXTRACTEURS:
        return jsonify({
            "error": f"Format non supporté : {extension}",
            "formats_supportes": list(EXTRACTEURS.keys())
        }), 415

    contenu = fichier.read()

    if not contenu:
        return jsonify({"error": "Fichier vide."}), 400

    try:
        extracteur = EXTRACTEURS[extension]
        texte_brut = extracteur(contenu)
        texte = nettoyer_texte(texte_brut)
        langue = detecter_langue(texte)

        return jsonify({
            "texte": texte,
            "langue": langue,
            "format": extension,
            "caracteres": len(texte),
            "succes": True
        })

    except Exception as e:
        return jsonify({
            "error": str(e),
            "format": extension,
            "succes": False
        }), 500


# ── Route de santé ────────────────────────────────────────────────────────
@app.route("/health", methods=["GET"])
def health():
    try:
        tesseract_version = pytesseract.get_tesseract_version().__str__()
    except Exception:
        tesseract_version = "introuvable"

    return jsonify({
        "status": "ok",
        "service": "SICOT OCR Service",
        "formats_supportes": list(EXTRACTEURS.keys()),
        "tesseract": tesseract_version,
    })


# ── Démarrage ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print(f"✅ SICOT OCR Service démarré sur http://localhost:{PORT}")
    print(f"📋 Tesseract : {TESSERACT_CMD}")
    print(f"📋 Formats supportés : {', '.join(EXTRACTEURS.keys())}")
    serve(app, host="127.0.0.1", port=PORT)